"""Unified file reader — auto-detects format (PDF, Word, Excel, text)."""

from __future__ import annotations

import os
from typing import Any

from langchain.tools import tool

from agent.config import settings
from agent.tools.registry import register
from agent.utils.path_guard import safe_resolve

_MAX_FILE_SIZE = 50 * 1024 * 1024
_MAX_PAGES = 200
_MAX_ROWS = 5000

_DOC_EXTENSIONS = {".pdf", ".docx", ".xlsx"}
_TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml",
    ".html", ".css", ".xml", ".csv", ".log", ".ini", ".cfg",
    ".rst", ".toml", ".env", ".bat", ".sh", ".sql", ".java",
    ".cpp", ".c", ".h", ".go", ".rs", ".vue", ".jsx", ".tsx",
}


def _resolve_sandbox(filename: str) -> str:
    return safe_resolve(filename, settings.FILE_READER_ROOT)


def _check_size(path: str) -> str | None:
    size = os.path.getsize(path)
    if size > _MAX_FILE_SIZE:
        return f"文件过大（{size / 1024 / 1024:.1f}MB > 50MB），拒绝读取"
    return None


def _read_pdf(path: str, filename: str, max_pages: int = 50) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "Error: pypdf not installed. Run: pip install pypdf"

    try:
        reader = PdfReader(path)
        total_pages = len(reader.pages)
        pages_to_read = min(max_pages, total_pages)
        parts: list[str] = []

        for i in range(pages_to_read):
            page = reader.pages[i]
            text = page.extract_text() or ""
            text = text.strip()
            if text:
                parts.append(f"--- 第 {i + 1} 页 ---\n{text}")

        result = f"📄 **PDF**: {filename}（共 {total_pages} 页，读取 {pages_to_read} 页）\n\n"
        result += "\n\n".join(parts)

        if total_pages > pages_to_read:
            result += f"\n\n（仅读取前 {pages_to_read} 页，共 {total_pages} 页）"

        if len(result) > 200_000:
            result = result[:200_000] + "\n\n（内容过长，已截断）"

        return result
    except Exception as exc:
        return f"PDF 读取失败：{exc}"


def _read_docx(path: str, filename: str) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    try:
        doc = Document(path)
        parts: list[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                style_name = para.style.name if para.style else ""
                if "Heading" in style_name:
                    parts.append(f"# {text}")
                else:
                    parts.append(text)

        for table in doc.tables:
            parts.append("\n[表格]")
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                parts.append(" | ".join(cells))
            parts.append("")

        result = f"📝 **Word**: {filename}\n\n" + "\n".join(parts)

        if len(result) > 200_000:
            result = result[:200_000] + "\n\n（内容过长，已截断）"

        return result
    except Exception as exc:
        return f"DOCX 读取失败：{exc}"


def _read_xlsx(path: str, filename: str, sheet_name: str | None = None, max_rows: int = 500) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    try:
        wb = load_workbook(path, read_only=True, data_only=True)

        if sheet_name:
            if sheet_name not in wb.sheetnames:
                return f"工作表 '{sheet_name}' 不存在。可用: {', '.join(wb.sheetnames)}"
            ws = wb[sheet_name]
        else:
            ws = wb.active

        rows_data: list[list[Any]] = []
        for row in ws.iter_rows(max_row=max_rows, values_only=True):
            rows_data.append([str(c) if c is not None else "" for c in row])

        if not rows_data:
            wb.close()
            return f"📊 **Excel**: {filename} - 工作表 '{ws.title}' 为空"

        header = rows_data[0]
        data_rows = rows_data[1:]
        total_rows = ws.max_row or 0

        lines = [
            f"📊 **Excel**: {filename} | 工作表: '{ws.title}' | "
            f"列数: {len(header)} | 读取: {len(data_rows)} 行"
        ]

        lines.append("\n**表头**: " + " | ".join(str(h) for h in header))

        if data_rows:
            lines.append("\n**数据预览** (前 20 行):")
            for i, row in enumerate(data_rows[:20], 1):
                lines.append(f"  {i}. " + " | ".join(row))

        remaining = total_rows - len(data_rows)
        if remaining > 0:
            lines.append(f"\n（共 {total_rows} 行，仅读取前 {len(data_rows)} 行）")

        wb.close()

        result = "\n".join(lines)
        if len(result) > 200_000:
            result = result[:200_000] + "\n\n（内容过长，已截断）"

        return result
    except Exception as exc:
        return f"XLSX 读取失败：{exc}"


def _read_text(path: str, filename: str, max_chars: int = 100_000) -> str:
    try:
        with open(path, "r", encoding="utf-8") as f:
            content = f.read(max_chars + 1)
        if len(content) > max_chars:
            return (
                f"📄 **文本**: {filename}\n\n"
                f"{content[:max_chars]}\n\n（内容过长，已截断，共 {len(content)} 字符）"
            )
        return f"📄 **文本**: {filename}\n\n{content}"
    except UnicodeDecodeError:
        try:
            with open(path, "r", encoding="gbk") as f:
                content = f.read(max_chars + 1)
            if len(content) > max_chars:
                return (
                    f"📄 **文本**: {filename}\n\n"
                    f"{content[:max_chars]}\n\n（内容过长，已截断）"
                )
            return f"📄 **文本**: {filename}\n\n{content}"
        except Exception as exc:
            return f"文本读取失败：{exc}"
    except Exception as exc:
        return f"文本读取失败：{exc}"


@register
@tool
def file_reader(filename: str, max_pages: int = 50, max_rows: int = 500) -> str:
    """读取任意类型文件内容（仅限项目目录内）。自动识别 PDF、Word、Excel、文本等格式。

    Args:
        filename: 文件路径，如 "docs/report.pdf"、"data/sales.xlsx"、"src/main.py"
        max_pages: PDF 最多读取页数，默认 50
        max_rows: Excel 最多读取行数，默认 500
    """
    try:
        safe = _resolve_sandbox(filename)
    except ValueError as exc:
        return f"路径错误：{exc}"

    if not os.path.isfile(safe):
        return f"文件不存在：{filename}"

    size_err = _check_size(safe)
    if size_err:
        return size_err

    ext = os.path.splitext(safe)[1].lower()

    if ext == ".pdf":
        return _read_pdf(safe, filename, max_pages)
    elif ext == ".docx":
        return _read_docx(safe, filename)
    elif ext == ".xlsx":
        return _read_xlsx(safe, filename, max_rows=max_rows)
    elif ext in _TEXT_EXTENSIONS or not ext:
        return _read_text(safe, filename)
    else:
        return _read_text(safe, filename)